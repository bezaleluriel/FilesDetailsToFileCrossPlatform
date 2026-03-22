import os
import html
import datetime as _dt
from dataclasses import dataclass
from pathlib import Path


RUN_EXTS = {".exe", ".dll"}


@dataclass(frozen=True)
class FileRow:
    file_name: str
    last_2_paths: str
    inner_path: str
    run_v: str
    data_v: str
    created: str
    size_bytes: int
    full_path: str


def _format_dt_local(ts: float) -> str:
    dt = _dt.datetime.fromtimestamp(ts)
    return dt.strftime("%Y-%m-%d %H:%M:%S")


def _creation_time(path: str) -> float:
    st = os.stat(path)
    # macOS exposes st_birthtime; Windows creation time is st_ctime.
    if hasattr(st, "st_birthtime"):
        return float(getattr(st, "st_birthtime"))
    return float(st.st_ctime)


def _last_two_dirs_within_root(root: Path, rel_parent: Path) -> str:
    # Based on folders *inside* the selected root.
    parts = [p for p in rel_parent.parts if p not in (".", "")]
    if not parts:
        return root.name
    if len(parts) == 1:
        return parts[0]
    return f"{parts[-2]}/{parts[-1]}"


def scan_folder(root_dir: str) -> list[FileRow]:
    root = Path(root_dir).resolve()
    rows: list[FileRow] = []

    for dirpath, _dirnames, filenames in os.walk(root):
        for fn in filenames:
            full = Path(dirpath) / fn
            try:
                st = os.stat(full)
            except OSError:
                continue

            ext = full.suffix.lower()
            is_run = ext in RUN_EXTS
            created_ts = _creation_time(str(full))

            parent = full.parent
            try:
                rel_parent = parent.resolve().relative_to(root)
            except Exception:
                rel_parent = Path(os.path.relpath(str(parent), str(root)))
            rel_parent_str = rel_parent.as_posix()
            inner_path = "" if rel_parent_str in (".", "") else rel_parent_str

            rows.append(
                FileRow(
                    file_name=full.name,
                    last_2_paths=_last_two_dirs_within_root(root, rel_parent),
                    inner_path=inner_path,
                    run_v="V" if is_run else "",
                    data_v="" if is_run else "V",
                    created=_format_dt_local(created_ts),
                    size_bytes=int(st.st_size),
                    full_path=str(full),
                )
            )

    rows.sort(key=lambda r: (r.inner_path, r.file_name.lower()))
    return rows


HEAD_COLS = [
    "File name",
    "Last 2 Paths",
    "Inner path",
    "Run (exe/dll)",
    "Data (other)",
    "Created (local)",
    "Size (bytes)",
]


def rows_to_tsv(rows: list[FileRow]) -> str:
    out_lines = ["\t".join(HEAD_COLS)]
    for r in rows:
        out_lines.append(
            "\t".join(
                [
                    r.file_name,
                    r.last_2_paths,
                    r.inner_path,
                    r.run_v,
                    r.data_v,
                    r.created,
                    str(r.size_bytes),
                ]
            )
        )
    return "\n".join(out_lines)


def rows_to_html_table(rows: list[FileRow]) -> str:
    def td(val: str) -> str:
        return f"<td>{html.escape(val)}</td>"

    def tdn(val: int) -> str:
        return f"<td style='text-align:right'>{val}</td>"

    body_rows = []
    for r in rows:
        body_rows.append(
            "<tr>"
            + td(r.file_name)
            + td(r.last_2_paths)
            + td(r.inner_path)
            + td(r.run_v)
            + td(r.data_v)
            + td(r.created)
            + tdn(r.size_bytes)
            + "</tr>"
        )

    return (
        "<table>"
        "<thead>"
        f"<tr>{''.join(f'<th>{html.escape(c)}</th>' for c in HEAD_COLS)}</tr>"
        "</thead>"
        "<tbody>"
        + "".join(body_rows)
        + "</tbody>"
        "</table>"
    )


def rows_to_html_document(rows: list[FileRow], title: str) -> str:
    table_html = rows_to_html_table(rows)
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: -apple-system, Segoe UI, Arial, sans-serif; padding: 24px; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #d0d0d0; padding: 6px 8px; font-size: 12px; }}
    th {{ background: #f6f6f6; text-align: left; }}
  </style>
</head>
<body>
  <h2>{html.escape(title)}</h2>
  {table_html}
</body>
</html>
"""


def export_docx(rows: list[FileRow], out_path: str, title: str) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install -r requirements.txt"
        ) from e

    doc = Document()
    doc.add_heading(title, level=2)

    getters = [
        ("File name", lambda r: r.file_name),
        ("Last 2 Paths", lambda r: r.last_2_paths),
        ("Inner path", lambda r: r.inner_path),
        ("Run (exe/dll)", lambda r: r.run_v),
        ("Data (other)", lambda r: r.data_v),
        ("Created (local)", lambda r: r.created),
        ("Size (bytes)", lambda r: str(r.size_bytes)),
    ]

    table = doc.add_table(rows=1, cols=len(getters))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, (name, _getter) in enumerate(getters):
        hdr_cells[i].text = name

    for r in rows:
        row_cells = table.add_row().cells
        for i, (_name, getter) in enumerate(getters):
            row_cells[i].text = getter(r)

    doc.save(out_path)

